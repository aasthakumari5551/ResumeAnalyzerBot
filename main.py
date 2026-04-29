from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    MessageHandler,
    filters,
    ContextTypes,
    CallbackQueryHandler,
    CommandHandler
)

import pdfplumber
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

import google.generativeai as genai
from dotenv import load_dotenv

import os

# 🔥 Load environment variables
load_dotenv(dotenv_path=".env")

# 🔥 Telegram Token
TOKEN = os.getenv("BOT_TOKEN")

# 🔥 Gemini API Setup
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

model = genai.GenerativeModel("gemini-1.5-flash")

# 🔥 Store user data
user_data = {}

# 🚀 Start Command
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):

    await update.message.reply_text(
        "🤖 Welcome to AI Resume Analyzer Bot!\n\n"
        "📄 Upload your resume PDF to begin."
    )


# 🔥 Extract text from PDF
def extract_text(file_bytes):

    text = ""

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:

        for page in pdf.pages:

            t = page.extract_text()

            if t:
                text += t + "\n"

    return text


# 🔹 Extract Keywords
def extract_keywords(text):

    words = text.split()

    return list(
        set(
            [
                w.strip(",.:;").lower()
                for w in words
                if len(w) > 3
            ]
        )
    )


# 🔥 ATS Score Logic
# 🔥 ATS Score Logic
def ats_score(resume, jd):

    important_skills = [

        # Programming Languages
        "python", "java", "c", "c++", "javascript",
        "typescript", "php", "ruby", "golang",
        "swift", "kotlin", "sql",

        # Web Development
        "html", "css", "react", "nextjs",
        "nodejs", "express", "fastapi",
        "django", "flask", "rest api",
        "graphql", "bootstrap", "tailwind",

        # Database
        "mysql", "postgresql", "mongodb",
        "firebase", "oracle", "sqlite",

        # Cloud / DevOps
        "aws", "azure", "gcp", "docker",
        "kubernetes", "jenkins", "linux",
        "github", "git", "ci/cd",

        # AI / ML / Data
        "machine learning", "deep learning",
        "artificial intelligence", "nlp",
        "tensorflow", "pytorch", "opencv",
        "pandas", "numpy", "data analysis",
        "prompt engineering", "llm",

        # Mechanical Engineering
        "mechanical", "manufacturing",
        "solidworks", "autocad", "cad",
        "catia", "ansys", "design",
        "simulation", "thermodynamics",
        "automation", "production",

        # Electrical / Electronics
        "embedded systems", "vlsi",
        "pcb", "microcontroller",
        "arduino", "raspberry pi",

        # Product Support / Support Roles
        "customer support", "troubleshooting",
        "ticketing", "sla", "technical support",
        "problem-solving", "analytical",
        "communication", "documentation",

        # Office / Business
        "excel", "powerpoint", "word",
        "data entry", "email handling",
        "reporting", "presentation",

        # Soft Skills
        "leadership", "teamwork",
        "collaboration", "adaptability",
        "critical thinking", "time management",

        # Cybersecurity
        "network security", "ethical hacking",
        "penetration testing", "cybersecurity",

        # Mobile Development
        "android", "ios", "flutter",
        "react native",

        # Testing
        "testing", "selenium",
        "unit testing", "debugging"
    ]

    resume_lower = resume.lower()

    jd_lower = jd.lower()

    matched_skills = []

    missing_skills = []

    for skill in important_skills:

        if skill in jd_lower:

            if skill in resume_lower:

                matched_skills.append(skill)

            else:

                missing_skills.append(skill)

    total_required = (
        len(matched_skills)
        + len(missing_skills)
    )

    if total_required == 0:

        score = 0

    else:

        score = round(
            (len(matched_skills) / total_required) * 10
        )

    # Bonus for projects or experience
    if (
        "project" in resume_lower
        or "experience" in resume_lower
    ):

        score += 1

    # Max score limit
    score = min(score, 10)

    # Feedback
    if score >= 8:

        feedback = "✅ Excellent ATS Match"

    elif score >= 6:

        feedback = "✅ Good ATS Match"

    elif score >= 4:

        feedback = "⚠️ Average ATS Match"

    else:

        feedback = "❌ Needs Improvement"

    return (
        score,
        matched_skills,
        missing_skills,
        feedback
    )

# 🔥 Gemini AI Function
def ask_ai(resume, jd, question):

    try:

        prompt = f"""
You are an expert ATS Resume Analyzer.

Analyze the resume according to the job description.

Resume:
{resume}

Job Description:
{jd}

Question:
{question}
"""

        response = model.generate_content(prompt)

        return response.text

    except Exception as e:

        return f"⚠️ AI error: {str(e)}"


# 🔥 Enhance Resume
# 🔥 Enhance Resume
def enhance_resume(resume_text, jd_text):

    important_skills = [

        "python", "java", "sql", "aws",
        "react", "excel", "communication",
        "customer support", "documentation",
        "problem-solving", "analytical",
        "cad", "autocad", "solidworks",
        "machine learning", "ai",
        "data analysis", "testing",
        "debugging", "automation"
    ]

    resume_lower = resume_text.lower()

    jd_lower = jd_text.lower()

    suggested_skills = []

    for skill in important_skills:

        if (
            skill in jd_lower
            and skill not in resume_lower
        ):

            suggested_skills.append(skill)

    lines = resume_text.split("\n")

    enhanced = []

    skills_added = False

    for line in lines:

        enhanced.append(line)

        # Add recommended skills only once
        if (
    line.strip().lower() == "skills"
    and not skills_added
):
            if suggested_skills:

                enhanced.append(
                    "- Recommended Skills: "
                    + ", ".join(suggested_skills[:8])
                )

            skills_added = True

    return "\n".join(enhanced)

# 📄 DOCX Template
def template_professional(text):

    doc = Document()

    doc.add_heading("RESUME", 0)

    for line in text.split("\n"):

        if not line.strip():
            continue

        if line.isupper():

            doc.add_heading(line, level=1)

        else:

            doc.add_paragraph(line)

    return doc


# 📄 PDF Generator
def create_pdf(text, filename):

    pdf = canvas.Canvas(filename, pagesize=A4)

    width, height = A4

    x = 40
    y = height - 40

    for line in text.split("\n"):

        line = line.strip()

        if not line:

            y -= 10
            continue

        while len(line) > 90:

            pdf.drawString(x, y, line[:90])

            line = line[90:]

            y -= 15

        pdf.drawString(x, y, line)

        y -= 15

        if y < 50:

            pdf.showPage()

            y = height - 40

    pdf.save()


# 🔘 Interactive Buttons
def generate_button():

    keyboard = [
        [
            InlineKeyboardButton(
                "📄 Generate Resume",
                callback_data="generate_resume"
            )
        ],
        [
            InlineKeyboardButton(
                "📊 Missing Skills",
                callback_data="missing_skills"
            )
        ],
        [
            InlineKeyboardButton(
                "📜 Certifications",
                callback_data="certs"
            )
        ],
        [
            InlineKeyboardButton(
                "💡 Improve Resume",
                callback_data="improve"
            )
        ]
    ]

    return InlineKeyboardMarkup(keyboard)


# 📄 Handle Resume Upload
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):

    file = await update.message.document.get_file()

    file_bytes = await file.download_as_bytearray()

    text = extract_text(file_bytes)

    user_data[update.effective_chat.id] = {
        "resume": text,
        "stage": "jd"
    }

    await update.message.reply_text(
        "✅ Resume uploaded!\n\n"
        "📄 Now send Job Description."
    )


# 🧠 Handle Text
async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):

    chat_id = update.effective_chat.id

    text = update.message.text

    if chat_id not in user_data:

        await update.message.reply_text(
            "❌ Upload resume first."
        )

        return

    stage = user_data[chat_id]["stage"]

    # 📊 JD Analysis Step
    if stage == "jd":

        resume = user_data[chat_id]["resume"]

        jd = text

        score, matched, missing, feedback = ats_score(
            resume,
            jd
        )

        user_data[chat_id]["analysis"] = {
            "score": score,
            "matched": list(matched),
            "missing": list(missing),
            "jd": jd
        }

        user_data[chat_id]["jd"] = jd

        user_data[chat_id]["stage"] = "qa"

        await update.message.reply_text(
            f"📊 Score: {score}/10\n"
            f"✅ Matched: {len(matched)}\n"
            f"❌ Missing: {len(missing)}\n\n"
            f"{feedback}"
        )

        await update.message.reply_text(
            "🤖 I'm your Resume Assistant!\n\n"
            "You can ask:\n"
            "• What skills are missing?\n"
            "• Am I fit for this role?\n"
            "• How to improve resume?\n"
            "• Suggest certifications\n\n"
            "Or use buttons 👇",
            reply_markup=generate_button()
        )

        return

    # 💬 AI Chat Mode
    if stage == "qa":

        resume = user_data[chat_id]["resume"]

        jd = user_data[chat_id]["jd"]

        await update.message.reply_text(
            "🤖 Thinking..."
        )

        answer = ask_ai(resume, jd, text)

        await update.message.reply_text(answer)

        await update.message.reply_text(
            "👇 Try more options:",
            reply_markup=generate_button()
        )

        return


# 🔘 Button Handler
async def handle_buttons(update: Update, context: ContextTypes.DEFAULT_TYPE):

    query = update.callback_query

    await query.answer()

    chat_id = query.message.chat.id

    data = query.data

    resume = user_data[chat_id]["resume"]

    jd = user_data[chat_id]["jd"]

    # 📄 Generate Resume
    if data == "generate_resume":

        enhanced = enhance_resume(resume, jd)

        doc = template_professional(enhanced)

        doc_name = f"resume_{chat_id}.docx"

        pdf_name = f"resume_{chat_id}.pdf"

        doc.save(doc_name)

        create_pdf(enhanced, pdf_name)

        await query.message.reply_text(
            "✅ Generating resume..."
        )

        await query.message.reply_document(
            open(doc_name, "rb")
        )

        await query.message.reply_document(
            open(pdf_name, "rb")
        )

    # 📊 Missing Skills
    elif data == "missing_skills":

        missing = user_data[chat_id]["analysis"]["missing"]

        await query.message.reply_text(
            "🚀 Missing Skills:\n"
            + ", ".join(list(missing)[:15])
        )

    # 📜 Certifications
    elif data == "certs":

        await query.message.reply_text(
            "📜 Recommended Certifications:\n"
            "AWS, Google Cloud, Azure, DSA, ML"
        )

    # 💡 Improve Resume
    elif data == "improve":

        await query.message.reply_text(
            "💡 Tips:\n"
            "- Add achievements\n"
            "- Use action verbs\n"
            "- Align with JD\n"
            "- Highlight projects"
        )


# 🚀 Run Bot
app = ApplicationBuilder().token(TOKEN).build()

# Start Command
app.add_handler(
    CommandHandler("start", start)
)

# Resume Upload
app.add_handler(
    MessageHandler(
        filters.Document.ALL,
        handle_file
    )
)

# Text Messages
app.add_handler(
    MessageHandler(
        filters.TEXT & ~filters.COMMAND,
        handle_text
    )
)

# Button Actions
app.add_handler(
    CallbackQueryHandler(handle_buttons)
)

print("🤖 Interactive AI Resume Bot running...")

app.run_polling()